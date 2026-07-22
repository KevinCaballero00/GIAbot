"""
Lector de documentos subidos por el docente para llenar las secciones 2/3/4
del FO-IN-17 (trabajos de grado, eventos, otras actividades) sin tener que
digitarlas paso a paso por chat.

Soporta `.txt`, `.docx` y `.pdf`. Las dependencias de `.docx`/`.pdf` se
importan de forma perezosa: si no están instaladas, `extraer_texto` lanza un
`ValueError` con un mensaje claro en español (el flujo de chat puede entonces
ofrecer volver al modo paso a paso), pero `.txt` sigue funcionando siempre.

API pública: extraer_texto(nombre_archivo, contenido_bytes) -> str
"""
from __future__ import annotations

MAX_BYTES = 2 * 1024 * 1024  # 2 MB
MAX_CHARS = 40000  # mismo tope que usa el estructurador para el texto crudo


def _extraer_txt(contenido_bytes: bytes) -> str:
    try:
        return contenido_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return contenido_bytes.decode("latin-1")


def _extraer_docx(contenido_bytes: bytes) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ValueError(
            "El servidor no tiene instalada la dependencia para leer archivos .docx "
            "(python-docx). Usa un archivo .txt o completa la información paso a paso."
        ) from exc

    import io
    documento = docx.Document(io.BytesIO(contenido_bytes))
    partes = [p.text for p in documento.paragraphs if p.text.strip()]
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _extraer_pdf(contenido_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ValueError(
            "El servidor no tiene instalada la dependencia para leer archivos .pdf "
            "(pdfplumber). Usa un archivo .txt/.docx o completa la información paso a paso."
        ) from exc

    import io
    partes: list[str] = []
    with pdfplumber.open(io.BytesIO(contenido_bytes)) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text() or ""
            if texto.strip():
                partes.append(texto)
    return "\n\n".join(partes)


_EXTRACTORES = {
    "txt": _extraer_txt,
    "docx": _extraer_docx,
    "pdf": _extraer_pdf,
}


def extraer_texto(nombre_archivo: str, contenido_bytes: bytes) -> str:
    """
    Extrae el texto plano de un archivo `.txt`, `.docx` o `.pdf` subido por el
    docente. Lanza `ValueError` (mensaje en español) si el archivo está vacío,
    supera `MAX_BYTES`, tiene una extensión no soportada, o falta la
    dependencia necesaria para leerlo.
    """
    if not contenido_bytes:
        raise ValueError("El archivo está vacío.")

    if len(contenido_bytes) > MAX_BYTES:
        raise ValueError(
            f"El archivo supera el tamaño máximo permitido ({MAX_BYTES // (1024 * 1024)} MB)."
        )

    extension = (nombre_archivo or "").rsplit(".", 1)[-1].lower() if "." in (nombre_archivo or "") else ""
    extractor = _EXTRACTORES.get(extension)
    if extractor is None:
        raise ValueError(
            f"Extensión '.{extension}' no soportada. Sube un archivo .txt, .docx o .pdf."
        )

    texto = extractor(contenido_bytes).strip()
    if not texto:
        raise ValueError("No se pudo extraer texto del archivo (¿está vacío o es una imagen escaneada?).")

    return texto[:MAX_CHARS]
